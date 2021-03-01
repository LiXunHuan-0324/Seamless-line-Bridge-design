# -*- coding: utf-8 -*-

'''
#  程序名:轨道强度检算
#  作  者:李建英
#         中南大学土木工程学院
#           2020年3月13日06:21:04
#
  说   明:铁路轨道静力计算
#  功能说明：
#  修改备注:
#  环境：Python3.8 win7
#  依赖库：
# 1 docx：请从https://www.lfd.uci.edu/~gohlke/pythonlibs/下载最新docx
# 2 word2010以上
'''

from math import exp
from math import sin
from math import cos
from math import tan
from math import sqrt
from math import pi
from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor, Cm
import judge
import os


class constant0():
    # 用于数据，坚守函数间传递的变量数量
    pass


CC = constant0()


def read_file(inname):
    """
    从文件对原始数据
    输入：
    inname：存放原始数据的文件名，该文件为文本文件
    """
    item_name = None
    # with open(inname,'r', encoding='utf-8', errors='ignore') as fin:

    with open(inname, 'r', encoding='utf-8') as fin:
        for oneline in fin:
            oneline = oneline.lstrip().rstrip()  # 去掉头尾的空格
            if oneline[:1] == "":
                break  # 空行 不处理
            if oneline[:1] == "#":
                continue  # 注释行 不处理
            if oneline[:1] == "[":  # 字段名
                item_name = oneline[1:-1]
                continue
            if item_name is None:
                continue
            aa = oneline.split(',')
            aa = [a.lstrip().rstrip() for a in aa]  # 去左右空格

            if item_name == 'LOC':
                CC.LOC_type = aa[0]  # 机车类型
            elif item_name == 'RAIL':
                CC.Rail_type = int(aa[0])  # 钢轨类型
                CC.Rail_mat = int(aa[1])  # 钢轨材料
                CC.Rail_length = float(aa[2])  # 钢轨长度
                CC.Rail_wear = float(aa[3])  # 钢轨磨耗量
            elif item_name == 'SLEEPER':
                CC.S_type = int(aa[0])  # 轨枕类型
                CC.S_number = float(aa[1])  # 轨枕根数
            elif item_name == 'BALLAST':
                CC.B_type = int(aa[0])  # 道床类型
                CC.H_thickness1 = float(aa[1])  # 道床面砟厚度
                CC.H_thickness2 = float(aa[2])  # 道床面砟厚度
            elif item_name == 'SUBGRADE':
                CC.L_type = int(aa[0])  # 线路类型
            elif item_name == 'RADIUS':
                CC.R_railway = int(aa[0])  # 曲线半径
            elif item_name == 'JUDGE':
                CC.r_Rail = float(aa[0])
                CC.r0 = float(aa[1])
                CC.brige_arclenth = float(aa[2])
                str_left_type = list(aa[3])
                CC.left_type = []
                for key in str_left_type:
                    CC.left_type.append(int(key))
                print(CC.left_type)
                CC.num_arc = len(CC.left_type)
                print("CC.num_arc#", CC.num_arc)
                CC.l0 = float(aa[4])
            else:
                print(" 字段错误 " + str(item_name))
                exit(0)
    print("输入完成,开始计算...")


def creat_constant():
    "根据输入数据生成计算中需要的其他参数"
    # 机车
    CC.Title = ''
    if CC.LOC_type == 'DF4':
        CC.Title = '东风四内燃机车'
        CC.P = [112800, 112800, 112800]  # N DF4轮重
        CC.x = [0.0, 1800.0, 3600.0]  # 车轮坐标
        CC.Train_speed = 120  # 直线速度
        CC.nwheel = len(CC.P)
        CC.i = range(CC.nwheel)  # 车轮编号
        CC.njisuan = [0, 1]  # 计算轮位编号
    elif CC.LOC_type == 'SS8':
        CC.Title = '韶山八电力机车'
        CC.P = [110000, 110000]  # N SS8轮重
        CC.x = [0.0, 2900.0]  # 车轮坐标
        CC.Train_speed = 160  # 直线速度
        CC.nwheel = len(CC.P)
        CC.i = range(CC.nwheel)  # 车轮编号
        CC.njisuan = [0, ]  # 计算轮位编号
    else:
        print(" 暂没考虑这种机车类型 " + str(CC.LOC_type))
        raise
    # 钢轨
    CC.E_rail = 2.1e5  # 钢轨弹性模量 N/mm**2
    if CC.Rail_mat == 1:
        CC.Rail_mat_name = "普通碳素钢"
        CC.Yield_strength = 405  # 钢轨屈服强度，与钢号有关 MPa
    else:
        print(' 暂未考虑这种钢轨材料 ' + str(CC.Rail_mat))
        raise
    if CC.Rail_type == 60:
        CC.J_Rail = 10.48
        #CC.F_rail = 77.45
        CC.F_rail = 65.8
        if CC.Rail_wear == 0:
            CC.J_rail = 3217.0e4  # 钢轨绕水平轴惯性矩 mm**4 与钢轨类型和磨耗有关
            CC.W_rail1 = 396.0e3  # 轨底断面系数  mm**3  与钢轨类型有关
            CC.W_rail2 = 339.4e3  # 轨头断面系数  mm**3  与钢轨类型有关
        else:
            print(' 暂未考虑 ' + str(CC.Rail_type) +
                  '   的这种磨耗量 ' + str(CC.Rail_wear))
            raise
    else:
        print(' 暂未考虑这种钢轨类型 ' + str(CC.Rail_type))
        raise
    # 轨枕
    CC.Asleeper = 1000000 / CC.S_number  # 计算轨枕间距
    # 道床
    if CC.B_type == 1:
        CC.sigma_h_Allowable = 0.5
    elif CC.B_type == 2:
        CC.sigma_h_Allowable = 0.4
    elif CC.B_type == 3:
        CC.sigma_h_Allowable = 0.3
    else:
        print("error in m ")
        raise
    # 路基
    if CC.L_type == 1:
        CC.L_type_name = "新建线路"
        CC.sigma_r_Allowable = 0.13
    elif CC.L_type == 2:
        CC.L_type_name = "既有线路"
        CC.sigma_r_Allowable = 0.15
    else:
        print("error in constant.L_type ")
        raise
    # 线路半径
    if CC.R_railway > 0:  # 曲线
        CC.Title = CC.Title + '曲线地段'
    else:
        CC.Title = CC.Title + '直线地段'

    CC.op_rail = 1.4e-5  # 轨道塑性原始弯曲曲率 cm**-1
    CC.oe_rail = 3.6e-7  # 轨道弹性原始弯曲曲率 cm**-1
    CC.beta_rail = 1.0  # 轨道刚度框架系数
    CC.f_rail = 0.2  # 轨道弯曲变形矢度 cm
    CC.Kl = 1.3  # 安全系数
    CC.alpha = 1.18e-5  # α，钢轨钢线性膨胀系数 /℃
    CC.Q_B = 84.3

    CC.Title = CC.Title + '轨道强度检算'
    CC.D_track1 = 30000.0  # 检算钢轨强度时，钢轨支点弹性系数 N/mm
    CC.D_track2 = 70000.0  # 检算轨下基础时，钢轨支点弹性系数 N/mm
    CC.m = 1.6  # 道床应力分布不均匀系数

    CC.u1 = CC.D_track1 / CC.Asleeper  # 计算弯矩
    CC.k1 = pow(CC.u1 / (4.0 * CC.E_rail * CC.J_rail), 1 / 4.0)  # 计算弯矩
    CC.u2 = CC.D_track2 / CC.Asleeper  # 计算反力
    CC.k2 = pow(CC.u2 / (4.0 * CC.E_rail * CC.J_rail), 1 / 4.0)  # 计算反力

    if CC.R_railway > 0:  # 曲线 注意若是直线 这参数为0
        v = 4.3 * sqrt(float(CC.R_railway))
        if CC.Train_speed > v:
            CC.Train_speed = v

    # 速度系数
    if CC.LOC_type == 'DF4':
        if CC.Train_speed > 160.0:
            CC.Alpha = 1.0
        else:
            CC.Alpha = 0.004 * CC.Train_speed
    elif CC.LOC_type == 'SS8':
        if CC.Train_speed > 160.0:
            CC.Alpha = 1.0
        else:
            CC.Alpha = 0.006 * CC.Train_speed
    else:
        print(" 暂没考虑这种机车类型 " + str(CC.LOC_type))
        raise
    # 偏载系数
    if CC.R_railway == 0:  # 直线时取4mm
        Delta_h = 4
    else:
        Delta_h = 75
    CC.Beta = 0.002 * Delta_h
    # 横向水平力系数
    if CC.R_railway == 0:
        CC.f = 1.25
    elif CC.R_railway >= 2000:
        CC.f = 1.3
    elif CC.R_railway >= 800 and CC.R_railway < 2000:
        CC.f = 1.45
    elif CC.R_railway >= 600 and CC.R_railway < 800:
        CC.f = 1.6
    elif CC.R_railway >= 500 and CC.R_railway < 600:
        CC.f = 1.7
    elif CC.R_railway >= 400 and CC.R_railway < 600:
        CC.f = 1.8
    elif CC.R_railway >= 300 and CC.R_railway < 400:
        CC.f = 1.6
    else:
        print(" 曲线半径太小 获取横向水平力系数错误：" + str(R_railway))
        raise


def main(inname, ouname):
    read_file(inname)
    creat_constant()

    document = Document()  # 使用 docx.Document新建一个Word文件，并将数据保存至该文件中
    document.styles['Normal'].font.name = u'宋体'  # 设置默认字体
    document.styles["Normal"].font.size = Pt(10)
    document.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')

    printconstant(document)
    printhead2(document, "钢轨静力响应计算:")
    printtable0(document, "∑Pφ3")

    sumPfai3 = []
    for jisuan in CC.njisuan:
        sumPfai3.append(trackstrength0(jisuan, "rail", document))

    maxsumPfai3, maxsumPfai3_num = findmax_for_sumPfai31(
        sumPfai3)  # 从sumPfai3找到最大值，并记录最大值的编号

    trackstrength1(maxsumPfai3, maxsumPfai3_num, document)

    printhead2(document, "轨枕反力静力响应计算:")
    printtable0(document, "∑Pφ1")
    sumPfai1 = []
    for jisuan in CC.njisuan:
        sumPfai1.append(trackstrength0(jisuan, "sleeper", document))
    maxsumPfai1, maxsumPfai1_num = findmax_for_sumPfai31(
        sumPfai1)  # 从sumPfai3找到最大值，并记录最大值的编号

    sleeperstrength1(maxsumPfai1, maxsumPfai1_num, document)
    Tc_Allow(document)

    document.save(ouname + '.docx')


def Tc_Allow(document):

    Rpie = (1/(1/(100*CC.R_railway)+CC.op_rail))
    mid1 = CC.beta_rail * CC.E_rail*1.0e4 * CC.J_Rail * pi ** 2.0
    omega = mid1 * (CC.oe_rail + (4/(pi ** 3.0 * Rpie)))
    mid2 = (4 * CC.Q_B / pi ** 3) - ((omega * CC.oe_rail) / CC.f_rail)
    l_square = (omega + (omega ** 2.0 + mid2 * mid1 * CC.f_rail) ** 0.5) / mid2
    f_oe = l_square * CC.oe_rail
    P = (mid2 * (CC.f_rail + f_oe) / l_square + 4 * CC.Q_B * l_square / pi ** 3.0) / (
        CC.f_rail + f_oe + 4 * l_square / (pi ** 3.0 * Rpie))
    P_Allowable = round((P/CC.Kl), 2)
    detaTc_Allowable = P_Allowable / (2 * CC.E_rail * CC.alpha * CC.F_rail*100)

    printhead2(document, "稳定性验算:")
    table = document.add_table(2, 4, "Table Grid")
    table.alignment = WD_TABLE_ALIGNMENT.CENTER  # 表格居中
    table.cell(0, 1).text = "换算曲率 R'"
    table.cell(0, 2).text = '允许温度压力'
    table.cell(0, 3).text = '允许温升'
    table.cell(1, 0).text = '数值大小'
    table.cell(1, 1).text = str("{:.2f}".format(Rpie) + 'm**-1')
    table.cell(1, 2).text = str("{:e}".format(P_Allowable)+" N")
    table.cell(1, 3).text = str("{:.2f}".format(detaTc_Allowable) + " ℃")
    CC.delta_T = float(detaTc_Allowable)

def findmax_for_sumPfai31(onelist):
    '''
    找出列表中最大值及其地址
    '''
    m0 = max(onelist)
    i = onelist.index(m0)
    return m0, i


def printhead1(document, str0):
    "打印一级标题"
    head = document.add_heading('', level=1)  # 一级标题
    run = head.add_run(str0)
    head.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER  # 位置
    run.font.name = u'黑体'  # 字体
    run.bold = True
    run._element.rPr.rFonts.set(qn('w:eastAsia'), u'黑体')
    run.font.size = Pt(14)  # 字号
    run.font.color.rgb = RGBColor(0x00, 0x00, 0x00)  # 标题颜色


def printhead2(document, str0):
    "打印二级标题"
    head = document.add_heading('', level=2)
    head.paragraph_format.space_before = Pt(10)
    run = head.add_run(str0)
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x00, 0x00, 0x00)  # 标题颜色
    run.font.name = u'Arial'  # 字体1
    run._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')  # 汉字使用字体2->宋体


def printconstant(document):
    '''
    输出原始参数和其他计算关键参数
    :param document:
    :return:
    '''
    printhead1(document, CC.Title)
    printhead2(document, "原始计算数据:")

    table = document.add_table(5, 9, "Table Grid")
    table.alignment = WD_TABLE_ALIGNMENT.CENTER  # 表格居中
    col_width = {0: 1.69, 1: 1.69, 2: 1.11, 3: 2.27,
                 4: 1.69, 5: 2.04, 6: 1.35, 7: 2.25, 8: 0.9}
    for col_num in range(9):
        table.cell(0, col_num).width = Cm(col_width[col_num])
    table.cell(0, 0).text = "钢轨"
    table.cell(0, 1).text = "类型"
    table.cell(0, 2).text = str(CC.Rail_type)
    table.cell(0, 3).text = "材料类型"
    # table.cell(0, 4).text = str(CC.Rail_mat)+'\n'+str(CC.Rail_mat_name)
    table.cell(0, 4).text = str(CC.Rail_mat)
    table.cell(0, 5).text = "长度m"
    table.cell(0, 6).text = str(format(CC.Rail_length, ".1f"))
    table.cell(0, 7).text = "磨耗量mm"
    table.cell(0, 8).text = str(format(CC.Rail_wear, ".0f"))
    table.cell(1, 0).text = "轨枕"
    table.cell(1, 1).text = "类型"
    table.cell(1, 2).text = str(CC.S_type)
    table.cell(1, 3).text = "数量 根/km"
    table.cell(1, 4).text = str(format(CC.S_number, ".0f"))
    table.cell(2, 0).text = "道床"
    table.cell(2, 1).text = "类型"
    table.cell(2, 2).text = str(CC.B_type)
    table.cell(2, 3).text = "面砟厚mm"
    table.cell(2, 4).text = str(format(CC.H_thickness1, ".0f"))
    table.cell(2, 5).text = "底砟厚mm"
    table.cell(2, 6).text = str(format(CC.H_thickness2, ".0f"))
    table.cell(3, 0).text = "路基"
    table.cell(3, 1).text = "类型"
    # table.cell(3, 2).text = str(CC.L_type)+'\n'+str(CC.L_type_name)
    table.cell(3, 2).text = str(CC.L_type)
    table.cell(4, 0).text = "平面"
    table.cell(4, 1).text = "半径m"
    table.cell(4, 2).text = str(format(CC.R_railway, ".0f"))

    printhead2(document, "相关计算参数:")
    table = document.add_table(4, 3, "Table Grid")
    table.cell(0, 1).text = "检算钢轨强度用"
    table.cell(0, 2).text = "检算轨下基础用"
    table.cell(1, 0).text = "钢轨支点弹性系数(N/mm)"
    table.cell(1, 1).text = str(CC.D_track1)
    table.cell(1, 2).text = str(CC.D_track2)
    table.cell(2, 0).text = "钢轨支点弹性系数(N/mm)"
    table.cell(2, 1).text = str(format(CC.u1, ".2f"))
    table.cell(2, 2).text = str(format(CC.u2, ".2f"))
    table.cell(3, 0).text = "刚比系数(1/mm)"
    table.cell(3, 1).text = str(format(CC.k1, ".9f"))
    table.cell(3, 2).text = str(format(CC.k2, ".9f"))

    printhead2(document, "准静态计算参数:")
    table = document.add_table(2, 5, "Table Grid")
    table.cell(0, 0).text = "曲线半径(m)"
    table.cell(0, 1).text = "列车速度(km/h)"
    table.cell(0, 2).text = "速度系数"
    table.cell(0, 3).text = "偏载系数"
    table.cell(0, 4).text = "横向水平力系数"
    table.cell(1, 0).text = str(CC.R_railway)
    table.cell(1, 1).text = str(format(CC.Train_speed, ".2f"))
    table.cell(1, 2).text = str(format(CC.Alpha, ".4f"))
    table.cell(1, 3).text = str(CC.Beta)
    table.cell(1, 4).text = str(CC.f)


def tableframe0(document, str0):
    '''
    创建表格框架，写表格第一列
    输入：
    document：
    str0：表格最后一列的字符串，创建钢轨计算表时用"∑Pφ3" 创建轨下计算表用"∑Pφ1"
    '''
    "创建钢轨计算的表框架 写入第1列"
    ncol = 3 + CC.nwheel
    table = document.add_table(5, ncol, "Table Grid")
    cell_1 = table.cell(0, 0)
    cell_2 = table.cell(4, 0)
    cell_1.merge(cell_2)
    cell_3 = table.cell(0, ncol - 1)
    cell_4 = table.cell(4, ncol - 1)
    cell_3.merge(cell_4)

    table.cell(0, 1).text = "p(N)"
    table.cell(1, 1).text = "x(mm)"
    table.cell(2, 1).text = "kx"
    table.cell(3, 1).text = "φ3"
    table.cell(4, 1).text = str0
    return table


def trackstrength0(jisuan, conflag, document):
    '''对一个计算轮，进行钢轨静弯矩、轨枕压力计算 并写入word文档的表格
    输入：
    jisuan：计算轮
    #document:word文档对象
    conflag:字符，控制参数 计算钢轨静弯矩时"rail" 计算轨枕静压力时"sleeper"
    返回：
    sum_pfai3：∑pφ3
    '''
    if conflag == "rail":
        k = CC.k1
        fai = fai3  # 函数名
        str0 = "∑Pφ3"
    elif conflag == "sleeper":
        k = CC.k2
        fai = fai1  # 函数名
        str0 = "∑Pφ1"
    else:
        print(" 控制参数错误 " + conflag)
        raise

    ncol = 3 + CC.nwheel
    table = tableframe0(document, str0)  # 生成表格 写第一列

    table.cell(0, 0).text = str(jisuan)  # 写计算轮位

    printonelineintable(table, 0, CC.P)
    x0 = [abs(a - CC.x[jisuan]) for a in CC.x]  # 变量a自在本句有效
    printx0 = [format(a, '.0f') for a in x0]
    printonelineintable(table, 1, printx0)

    kx0 = [k * x0[j] for j in CC.i]
    printkx0 = [format(a, '.7f') for a in kx0]
    printonelineintable(table, 2, printkx0)

    afai03 = [fai(a) for a in kx0]  # 变量a自在本句有效 fai是变量 计算fai的函数名
    printafai03 = [format(a, '.7f') for a in afai03]
    printonelineintable(table, 3, printafai03)

    pfai03 = [CC.P[j] * afai03[j] for j in CC.i]
    printapfai03 = [format(a, '.2f') for a in pfai03]
    printonelineintable(table, 4, printapfai03)

    sum_pfai3 = sum(pfai03)
    table.cell(0, ncol - 1).text = str(format(sum_pfai3, '.2f'))
    return sum_pfai3


def trackstrength1(sum_pfai3, maxsumPfai3_num, document):
    # 对钢轨强度进行检算，并存入word
    dap = document.add_paragraph
    dap("\n最不利轮位为 " + str(maxsumPfai3_num) +
        " 最大∑pφ3为 " + str(format(sum_pfai3, '.2f')))

    Mj = sum_pfai3 / (4 * CC.k1)
    Md = Mj * (1 + CC.Alpha + CC.Beta) * CC.f
    dap("最大静弯矩:" + str(format(Mj, '.2f')) + "N·mm" +
        " 最大动弯矩:" + str(format(Md, '.2f')) + "N·mm")

    sigma_Allowable = (CC.Yield_strength / 1.3)
    sigma_1d = float(Md / CC.W_rail1)
    sigma_2d = float(Md / CC.W_rail2)

    table = document.add_table(4, 4, "Table Grid")
    table.cell(0, 1).text = "计算值"
    table.cell(0, 2).text = "允许值"
    table.cell(0, 3).text = "检算结论"
    table.cell(1, 0).text = "轨底拉应力MPa"
    table.cell(1, 1).text = str(format(sigma_1d, '.1f'))
    table.cell(1, 2).text = str(format(sigma_Allowable, '.1f'))
    flag1 = sigma_1d < sigma_Allowable
    if flag1:
        table.cell(1, 3).text = "合格"
    else:
        table.cell(1, 3).text = "不合格"
    table.cell(2, 0).text = "轨头压应力MPa"
    table.cell(2, 1).text = str(format(sigma_2d, '.1f'))
    table.cell(2, 2).text = str(format(sigma_Allowable, '.1f'))
    flag2 = sigma_1d < sigma_Allowable
    if flag2:
        table.cell(2, 3).text = "合格"
    else:
        table.cell(2, 3).text = "不合格"
    table.cell(3, 0).text = "检算结论"
    cell_1 = table.cell(3, 1)
    cell_2 = table.cell(3, 3)
    cell_1.merge(cell_2)
    if flag1 and flag2:
        table.cell(3, 1).text = "钢轨强度检算合格"
    else:
        table.cell(3, 1).text = "钢轨强度检算不合格"


def printonelineintable(table, n, one):
    '''
    将准静态计算的一个列表写入表格一行，从第2列开始
    table：表格
    n：行号
    one：要写入的列表
    '''
    for i, a in enumerate(one):
        table.cell(n, i + 2).text = str(a)


def printtable0(document, str0):
    '''
    创建钢轨计算的表头
    输入：
    document：
    str0：表格最后一列的字符串，创建钢轨计算表时用"∑Pφ3" 创建轨下计算表用"∑Pφ1"
    '''
    ncol = 3 + CC.nwheel
    table = document.add_table(1, ncol, "Table Grid")
    table.cell(0, 0).text = "计算轮位"
    table.cell(0, 1).text = "计算值"
    table.cell(0, ncol - 1).text = str0
    for i in CC.i:
        table.cell(0, 2 + i).text = "轮位" + str(i + 1)


def sleeperstrength1(sum_pfai1, maxsumPfai1_num, document):
    '''
    检算轨枕、道床、路基强度 将结论写出到word
    输入：
    sum_pfai1：最不利轮位的∑Pφ1
    maxsumPfai1_num：最不利轮位
    document：word对象
    '''
    dap = document.add_paragraph
    dap("\n最不利轮位为 " + str(maxsumPfai1_num) +
        " 最大∑pφ1为 " + str(format(sum_pfai1, '.2f')))

    Rj = sum_pfai1 * CC.k2 * CC.Asleeper / 2
    Rd = (1 + CC.Alpha + CC.Beta) * Rj
    dap("最大静压力:" + str(format(Rj, '.2f')) + "N" +
        " 最大动压力:" + str(format(Rd, '.2f')) + "N")
    # 轨枕强度检算
    a1, e, b0, Ks, L, Mg_Allowable, Mc_Allowable, b, e0 = sleeper()
    Mg = Ks * Rd * (a1 * a1 / (2 * e) - b0 / 8) / 1.0e6
    Mc = Ks * Rd * (3 * L ** 2 + 4 * e * e - 8 * a1 * e -
                    12 * a1 * L) / (12 * L + 8 * e) / 1.0e6

    # 道床强度检算
    sigma_b = CC.m * Rd / (b * e0)

    # 路基强度检算
    h1 = b / (2 * tan(35 * pi / 180))
    h2 = e0 / (2 * tan(35 * pi / 180))
    h = CC.H_thickness1 + CC.H_thickness2 / 2
    if h >= 0 and h <= h1:
        sigma_r = CC.m * Rd / (b * e0)
    elif h > h1 and h < h2:
        sigma_r = Rd / (2 * h * e0 * tan(35 * pi / 180))
    elif h > h2:
        sigma_r = Rd / (4 * h * h * tan(35 * pi / 180) * tan(35 * pi / 180))
    else:
        print(" 道床厚度错误 ")
        raise

    table = document.add_table(6, 4, "Table Grid")
    table.cell(0, 1).text = "计算值"
    table.cell(0, 2).text = "允许值"
    table.cell(0, 3).text = "检算结论"
    table.cell(1, 0).text = "轨下截面正弯矩MPa"
    table.cell(1, 1).text = str(format(Mg, '.1f'))
    table.cell(1, 2).text = str(format(Mg_Allowable, '.1f'))
    flag1 = Mg < Mg_Allowable
    if flag1:
        table.cell(1, 3).text = "合格"
    else:
        table.cell(1, 3).text = "不合格"
    table.cell(2, 0).text = "轨枕中间截面负弯矩"
    table.cell(2, 1).text = str(format(Mc, '.1f'))
    table.cell(2, 2).text = str(format(Mc_Allowable, '.1f'))
    flag2 = Mc < Mc_Allowable
    if flag2:
        table.cell(2, 3).text = "合格"
    else:
        table.cell(2, 3).text = "不合格"
    table.cell(3, 0).text = "道床顶面压应力MPa"
    table.cell(3, 1).text = str(format(sigma_b, '.2f'))
    table.cell(3, 2).text = str(format(CC.sigma_h_Allowable, '.2f'))
    flag3 = sigma_b < CC.sigma_h_Allowable
    if flag3:
        table.cell(3, 3).text = "合格"
    else:
        table.cell(3, 3).text = "不合格"
    table.cell(4, 0).text = "路基面应力MPa"
    table.cell(4, 1).text = str(format(sigma_r, '.2f'))
    table.cell(4, 2).text = str(format(CC.sigma_r_Allowable, '.2f'))
    flag4 = sigma_r < CC.sigma_r_Allowable
    if flag4:
        table.cell(4, 3).text = "合格"
    else:
        table.cell(4, 3).text = "不合格"
    table.cell(5, 0).text = "检算结论"
    cell_1 = table.cell(5, 1)
    cell_2 = table.cell(5, 3)
    cell_1.merge(cell_2)
    if flag1 and flag2 and flag3 and flag4:
        table.cell(5, 1).text = "轨下基础检算合格"
    else:
        table.cell(5, 1).text = "轨迹下基础检算不合格"


def fai3(kx):
    '''
    计算φ3，用于钢轨弯矩计算
    输入：
    kx：换算距离
    输出：
    fai3：φ3
    '''
    fai3 = exp(-kx) * (cos(kx) - sin(kx))
    return fai3


def fai1(kx):
    '''
    计算φ1，用于轨下基础计算
    输入：
    kx：换算距离
    输出：
    fai1：φ1
    '''
    fai1 = exp(-kx) * (cos(kx) + sin(kx))
    return fai1


def sleeper():
    '''取轨枕参数
    输入：
    m:道床的类型
    返回:
    sigma_h_Allowable:路基容许应力
    '''
    if CC.S_type == 1:
        a1 = 500
        e = 950
        b0 = 150
        Ks = 1
        L = 2500
        Mg_Allowable = 11.9
        Mc_Allowable = 8.8
        b = 275
        e0 = 950
    elif CC.S_type == 2:
        a1 = 500
        e = 950
        b0 = 150
        Ks = 1
        L = 2500
        Mg_Allowable = 13.3
        Mc_Allowable = 10.5
        b = 275
        e0 = 1175
    elif CC.S_type == 3:
        a1 = 500
        e = 950
        b0 = 150
        Ks = 1
        L = 2600
        Mg_Allowable = 18
        Mc_Allowable = 14
        b = 275
        e0 = 1175
    else:
        print(" error ")
    return (a1, e, b0, Ks, L, Mg_Allowable, Mc_Allowable, b, e0)


if __name__ == "__main__":
    print("\n开始运行轨道强度检算程序")
    print("\n请键入输入文件名（缺省为track.in）")
    inname = input()

    if inname == "":
        inname = 'track.in'

    basepath = os.path.abspath(__file__)
    folder = os.path.dirname(basepath)
    data_path = os.path.join(folder, inname)

    print("\n请键入输出文件名 不输入后缀（缺省为track）")
    ouname = input()
    if ouname == "":
        ouname = 'track'

    main(data_path, ouname)

    print("\n轨道强度检算结束 请检查文件 " + ouname + '.docx')
